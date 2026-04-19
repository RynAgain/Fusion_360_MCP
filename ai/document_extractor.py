"""Document text extraction for PDFs, DOCX, images, and text files.

Inspired by Roo Code's extract-text.ts. Provides structured text extraction
with metadata (line count, truncation, file type).
"""
import logging
import os
from pathlib import Path
from typing import Optional

logger = logging.getLogger(__name__)

# Supported file extensions
SUPPORTED_TEXT = {".txt", ".md", ".csv", ".tsv", ".json", ".xml", ".yaml", ".yml", ".ini", ".cfg", ".log"}
SUPPORTED_DOCS = {".pdf", ".docx"}
SUPPORTED_IMAGES = {".png", ".jpg", ".jpeg", ".gif", ".webp", ".bmp", ".tiff", ".svg"}
SUPPORTED_ALL = SUPPORTED_TEXT | SUPPORTED_DOCS | SUPPORTED_IMAGES

# Limits
MAX_FILE_SIZE_MB = 10
MAX_LINES = 2000
MAX_IMAGE_SIZE_MB = 5


def extract_text(file_path: str, max_lines: int = MAX_LINES) -> dict:
    """Extract text from a supported file format.

    Returns:
        dict with keys: content, total_lines, returned_lines, was_truncated,
        file_type, file_name, error (if any)
    """
    path = Path(file_path)

    if not path.exists():
        return {"error": f"File not found: {file_path}", "file_name": path.name}

    # Check file size
    size_mb = path.stat().st_size / (1024 * 1024)
    if size_mb > MAX_FILE_SIZE_MB:
        return {"error": f"File too large: {size_mb:.1f}MB (max {MAX_FILE_SIZE_MB}MB)", "file_name": path.name}

    ext = path.suffix.lower()

    if ext not in SUPPORTED_ALL:
        return {"error": f"Unsupported file type: {ext}", "file_name": path.name, "supported": list(SUPPORTED_ALL)}

    try:
        if ext in SUPPORTED_TEXT:
            return _extract_text_file(path, max_lines)
        elif ext == ".pdf":
            return _extract_pdf(path, max_lines)
        elif ext == ".docx":
            return _extract_docx(path, max_lines)
        elif ext in SUPPORTED_IMAGES:
            return _extract_image(path)
        else:
            return {"error": f"No handler for {ext}", "file_name": path.name}
    except Exception as exc:
        logger.exception("Failed to extract text from %s", file_path)
        return {"error": str(exc), "file_name": path.name}


def _extract_text_file(path: Path, max_lines: int) -> dict:
    """Extract text from plain text files."""
    try:
        text = path.read_text(encoding="utf-8", errors="replace")
    except UnicodeDecodeError:
        text = path.read_text(encoding="latin-1", errors="replace")

    lines = text.splitlines()
    total = len(lines)
    truncated = total > max_lines
    if truncated:
        lines = lines[:max_lines]

    return {
        "content": "\n".join(lines),
        "total_lines": total,
        "returned_lines": len(lines),
        "was_truncated": truncated,
        "file_type": path.suffix.lower(),
        "file_name": path.name,
    }


def _extract_pdf(path: Path, max_lines: int) -> dict:
    """Extract text from PDF using PyMuPDF (fitz) or pdfplumber."""
    text = ""

    # Try PyMuPDF first (faster, better layout)
    try:
        import fitz  # PyMuPDF
        doc = fitz.open(str(path))
        pages = []
        for page_num, page in enumerate(doc):
            page_text = page.get_text()
            if page_text.strip():
                pages.append(f"--- Page {page_num + 1} ---\n{page_text}")
        doc.close()
        text = "\n\n".join(pages)
    except ImportError:
        # Fallback to pdfplumber
        try:
            import pdfplumber
            with pdfplumber.open(str(path)) as pdf:
                pages = []
                for i, page in enumerate(pdf.pages):
                    page_text = page.extract_text() or ""
                    if page_text.strip():
                        pages.append(f"--- Page {i + 1} ---\n{page_text}")
                text = "\n\n".join(pages)
        except ImportError:
            return {
                "error": "PDF extraction requires 'pymupdf' or 'pdfplumber'. Install with: pip install pymupdf",
                "file_name": path.name,
            }

    if not text.strip():
        return {
            "content": "[PDF contains no extractable text -- may be scanned/image-based]",
            "total_lines": 0,
            "returned_lines": 0,
            "was_truncated": False,
            "file_type": ".pdf",
            "file_name": path.name,
        }

    lines = text.splitlines()
    total = len(lines)
    truncated = total > max_lines
    if truncated:
        lines = lines[:max_lines]

    return {
        "content": "\n".join(lines),
        "total_lines": total,
        "returned_lines": len(lines),
        "was_truncated": truncated,
        "file_type": ".pdf",
        "file_name": path.name,
    }


def _extract_docx(path: Path, max_lines: int) -> dict:
    """Extract text from DOCX files."""
    try:
        from docx import Document
    except ImportError:
        return {
            "error": "DOCX extraction requires 'python-docx'. Install with: pip install python-docx",
            "file_name": path.name,
        }

    doc = Document(str(path))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

    # Also extract from tables
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                paragraphs.append(" | ".join(cells))

    text = "\n".join(paragraphs)
    lines = text.splitlines()
    total = len(lines)
    truncated = total > max_lines
    if truncated:
        lines = lines[:max_lines]

    return {
        "content": "\n".join(lines),
        "total_lines": total,
        "returned_lines": len(lines),
        "was_truncated": truncated,
        "file_type": ".docx",
        "file_name": path.name,
    }


def _extract_image(path: Path) -> dict:
    """Process an image file -- return base64 for multimodal use."""
    import base64

    size_mb = path.stat().st_size / (1024 * 1024)
    if size_mb > MAX_IMAGE_SIZE_MB:
        return {
            "error": f"Image too large: {size_mb:.1f}MB (max {MAX_IMAGE_SIZE_MB}MB)",
            "file_name": path.name,
        }

    data = path.read_bytes()
    b64 = base64.b64encode(data).decode("ascii")

    # Determine media type
    ext = path.suffix.lower()
    media_types = {
        ".png": "image/png", ".jpg": "image/jpeg", ".jpeg": "image/jpeg",
        ".gif": "image/gif", ".webp": "image/webp", ".bmp": "image/bmp",
        ".tiff": "image/tiff", ".svg": "image/svg+xml",
    }
    media_type = media_types.get(ext, "application/octet-stream")

    return {
        "file_type": ext,
        "file_name": path.name,
        "media_type": media_type,
        "base64_data": b64,
        "size_bytes": len(data),
        "is_image": True,
    }


def get_supported_extensions() -> dict:
    """Return supported file extensions by category."""
    return {
        "text": sorted(SUPPORTED_TEXT),
        "documents": sorted(SUPPORTED_DOCS),
        "images": sorted(SUPPORTED_IMAGES),
    }
